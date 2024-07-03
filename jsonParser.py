# import json
# from bs4 import BeautifulSoup
# from mathml2docx import MathmlToDocx
# from docx import Document
# from docx.shared import Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement

# def process_html(html_content):
#     soup = BeautifulSoup(html_content, 'lxml')
#     for math_element in soup.find_all('span', class_='mathMlContainer'):
#         mathml = math_element.find('math')
#         if mathml:
#             mathml['xmlns'] = 'http://www.w3.org/1998/Math/MathML'
#             math_element.replace_with(mathml)
#     return str(soup)

# def set_two_columns(section):
#     sectPr = section._sectPr
#     cols = OxmlElement('w:cols')
#     cols.set(qn('w:num'), '2')
#     cols.set(qn('w:space'), '720')
#     sectPr.append(cols)

# def process_json_to_docx(json_file, output_file):
#     with open(json_file, 'r', encoding='utf-8') as f:
#         data = json.load(f)

#     questions = data['result']['data'][0]['sec_details'][0]['sec_questions']

#     document = Document()
#     converter = MathmlToDocx()

#     for section in document.sections:
#         set_two_columns(section)
#         section.top_margin = Inches(1)
#         section.bottom_margin = Inches(1)
#         section.left_margin = Inches(1)
#         section.right_margin = Inches(1)

#     for index, question in enumerate(questions, start=1):
#         heading = document.add_heading(f"Question {index}", level=1)
#         heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
#         q_string = process_html(question['que']['1']['q_string'])

#         # Add question text
#         converter.add_html_to_document(q_string, document)

#         if 'q_option' in question['que']['1']:
#             for i, option in enumerate(question['que']['1']['q_option'], start=97):
#                 option_text = process_html(option)

#                 # Add option text
#                 option_paragraph = document.add_paragraph(f"{chr(i)}. ", style='ListBullet')
#                 option_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
#                 converter.add_html_to_document(option_text, document)

#         document.add_paragraph()  # Add a blank line between questions

#     document.save(output_file)

# # File paths
# json_file = r'C:\Users\pavan\OneDrive\Desktop\Exampad\EXAMPAD\KODEMAPA-EXAMPAD\responseBody\Probability_Probability_(Hard).json'
# output_docx_file = 'output.docx'

# # Process JSON to DOCX
# process_json_to_docx(json_file, output_docx_file)

# print(f"DOCX file '{output_docx_file}' created successfully.")


# import json
# from bs4 import BeautifulSoup
# from mathml2docx import MathmlToDocx
# from docx import Document
# from docx.shared import Inches, Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement

# def process_html(html_content):
#     soup = BeautifulSoup(html_content, 'lxml')
#     for math_element in soup.find_all('span', class_='mathMlContainer'):
#         mathml = math_element.find('math')
#         if mathml:
#             mathml['xmlns'] = 'http://www.w3.org/1998/Math/MathML'
#             math_element.replace_with(mathml)
#     return str(soup)

# def set_two_columns(section):
#     sectPr = section._sectPr
#     cols = OxmlElement('w:cols')
#     cols.set(qn('w:num'), '2')
#     cols.set(qn('w:space'), '720')
#     sectPr.append(cols)

# def process_json_to_docx(json_file, output_file):
#     with open(json_file, 'r', encoding='utf-8') as f:
#         data = json.load(f)
    
#     questions = data['result']['data'][0]['sec_details'][0]['sec_questions']
    
#     document = Document()
#     converter = MathmlToDocx()
    
#     for section in document.sections:
#         set_two_columns(section)
#         section.top_margin = Inches(0.5)
#         section.bottom_margin = Inches(0.5)
#         section.left_margin = Inches(0.5)
#         section.right_margin = Inches(0.5)
    
#     for index, question in enumerate(questions, start=1):
#         # Add question number
#         heading = document.add_paragraph(f"Question {index}", style='Heading 1')
#         heading.runs[0].font.size = Pt(12)
#         heading.runs[0].font.bold = True
#         heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
#         # Add question text
#         q_string = process_html(question['que']['1']['q_string'])
#         q_paragraph = document.add_paragraph()
#         q_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
#         converter.add_html_to_document(q_string, document)
        
#         if 'q_option' in question['que']['1']:
#             for i, option in enumerate(question['que']['1']['q_option'], start=97):
#                 option_text = process_html(option)
                
#                 # Add option text
#                 option_paragraph = document.add_paragraph(style='List Bullet')
#                 option_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
#                 option_run = option_paragraph.add_run(f"{chr(i)}. ")
#                 option_run.font.bold = True
#                 converter.add_html_to_document(option_text, document)
        
#         document.add_paragraph()  # Add a blank line between questions
    
#     document.save(output_file)

# # File paths
# json_file = r'C:\Users\pavan\OneDrive\Desktop\Exampad\EXAMPAD\KODEMAPA-EXAMPAD\responseBody\Probability_Probability_(Hard).json'
# output_docx_file = 'output.docx'

# # Process JSON to DOCX
# process_json_to_docx(json_file, output_docx_file)

# print(f"DOCX file '{output_docx_file}' created successfully.")

#best
import json
from bs4 import BeautifulSoup
from mathml2docx import MathmlToDocx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def process_html(html_content):
    soup = BeautifulSoup(html_content, 'lxml')
    for math_element in soup.find_all('span', class_='mathMlContainer'):
        mathml = math_element.find('math')
        if mathml:
            mathml['xmlns'] = 'http://www.w3.org/1998/Math/MathML'
            math_element.replace_with(mathml)
    return str(soup)

def set_two_columns(section):
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')
    sectPr.append(cols)

def process_json_to_docx(json_file, output_file):
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    questions = data['result']['data'][0]['sec_details'][0]['sec_questions']
    
    document = Document()
    converter = MathmlToDocx()
    
    for section in document.sections:
        set_two_columns(section)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    for index, question in enumerate(questions, start=1):
        # Add question number
        heading = document.add_paragraph(f"Question {index}", style='Heading 1')
        heading.runs[0].font.size = Pt(12)
        heading.runs[0].font.bold = True
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add question text
        q_string = process_html(question['que']['1']['q_string'])
        converter.add_html_to_document(q_string, document)
        
        if 'q_option' in question['que']['1']:
            for i, option in enumerate(question['que']['1']['q_option'], start=97):
                option_text = process_html(option)
                
                # Add option text without the extra dot
                option_paragraph = document.add_paragraph(style='Normal')
                option_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                option_run = option_paragraph.add_run(f"{chr(i)}) ")
                option_run.font.bold = True
                converter.add_html_to_document(option_text, document)
    
    document.save(output_file)

# File paths
json_file = r'C:\Users\pavan\OneDrive\Desktop\Exampad\EXAMPAD\KODEMAPA-EXAMPAD\responseBody\Thermodynamics_Thermodynamics_(Difficult)_1.json'
output_docx_file = 'output.docx'

# Process JSON to DOCX
process_json_to_docx(json_file, output_docx_file)

print(f"DOCX file '{output_docx_file}' created successfully.")

