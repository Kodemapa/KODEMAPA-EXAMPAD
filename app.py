import ast
import html
import random
import re
from flask import Flask, jsonify, render_template, render_template_string, request, redirect, url_for , send_file
import json
import pypandoc
import requests
import urllib.parse
import subprocess
from docx import Document
from pdf2docx import Converter
from html import unescape
import os
import subprocess
from docx import Document
import webview
import tempfile
import pypandoc
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

pypandoc.download_pandoc()
app = Flask(__name__)

@app.route('/', methods=['GET', 'POST'])
def select_class():
    if request.method == 'POST':
        selected_class = request.form.get('class_selection')
        if selected_class:
            return redirect(url_for('l3_screen', class_selection=selected_class))
    
    
    return render_template('select_class.html')

@app.route('/l3_screen/<class_selection>', methods=['GET', 'POST'])
def l3_screen(class_selection):
    json_file = f'CBSE_{class_selection}.json'
    
    # Load data from selected JSON file
    with open(json_file, 'r') as file:
        data = json.load(file)['result']['data']
    
    if request.method == 'POST':
        selected_l3 = request.form.get('l3_category')
        if selected_l3:
            return redirect(url_for('l4_l5_screen', class_selection=class_selection, l3_id=selected_l3))
    
    return render_template('l3_screen.html', data=data)

@app.route('/l4_l5_screen/<class_selection>/<l3_id>', methods=['GET', 'POST'])
def l4_l5_screen(class_selection, l3_id):
    json_file = f'CBSE_{class_selection}.json'
    
    # Load data from selected JSON file
    with open(json_file, 'r') as file:
        data = json.load(file)['result']['data']
    
    l4_data = next((l3['L4'] for l3 in data['L3'] if str(l3['id']) == l3_id), [])
    
    if request.method == 'POST':
        selected_urls_by_category = {}
        
        for category in l4_data:
            selected_urls = request.form.getlist(category['name'])
            if selected_urls:
                selected_urls_by_category[category['name']] = ','.join(selected_urls)
        
        if selected_urls_by_category:
            # Serialize the dictionary to URL-encoded query string
            serialized_data = urllib.parse.urlencode(selected_urls_by_category)
            return redirect(url_for('test_page') + '?' + serialized_data)
    
    return render_template('l4_l5_screen.html', l4_data=l4_data)

@app.route('/test_page', methods=['GET', 'POST'])
def test_page():
    query_params = request.query_string.decode('utf-8')
    test_urls_by_category = urllib.parse.parse_qs(query_params)
    
    questions_by_section = {}
    selected_questions_ids = set()
    try:
        for category, test_urls in test_urls_by_category.items():
            questions_by_section[category] = []
            # Join the list of URLs back into a single string
            full_url_string = ','.join(test_urls)
            
            # Use regex to find all JSON-like structures
            json_strings = re.findall(r'\{[^}]+\}', full_url_string)
            
            for json_string in json_strings:
                try:
                    # Replace single quotes with double quotes and remove any leading/trailing '+'
                    json_string = json_string.replace("'", '"').strip('+')
                    test_url_dict = json.loads(json_string)
                except json.JSONDecodeError:
                    print(f"Error parsing URL for {category}: {json_string}")
                    continue  # Skip this problematic URL and continue with the next one

                json_file = test_url_dict.get("file_path")
                if not json_file:
                    print(f"No file_path found in: {test_url_dict}")
                    continue

                json_file = json_file.replace('\\', '/')
                # Load data from selected JSON file
                try:
                    with open(json_file, 'r') as file:
                        data = json.load(file)['result']['data']
                except FileNotFoundError:
                    print(f"File not found: {json_file}")
                    continue
                except json.JSONDecodeError:
                    print(f"Error decoding JSON file: {json_file}")
                    continue

                for test_info in data:
                    if 'sec_details' in test_info:
                        for sec_detail in test_info['sec_details']:
                            for question in sec_detail.get('sec_questions', []):
                                qid = question["qid"]
                                if qid not in selected_questions_ids:
                                    selected_questions_ids.add(qid)
                                    # Determine difficulty level if not provided
                                    if 'difficulty_level' not in question:
                                        question['difficulty_level'] = determine_difficulty(json_file)
                                    else:
                                        # Ensure difficulty_level is lowercase
                                        question['difficulty_level'] = question['difficulty_level'].lower()
                                    
                                    # Clean the q_string field
                                    q_string = question['que']['1']['q_string']
                                    
                                    # Remove content inside <b> tags
                                    q_string = re.sub(r'<b>.*?</b>', '', q_string)
                                    
                                    question['que']['1']['q_string'] = q_string

                                    questions_by_section[category].append(question)
                    else:
                        print(f"Warning: 'sec_details' not found in test_info: {test_info}")

    except Exception as e:
        return f"An error occurred: {str(e)}", 500
    
    if request.method == 'POST':
        selected_questions_indices = request.form.getlist('question')
        selected_questions = []
        selected_questions_ids=[]
        for idx in selected_questions_indices:
            section_name, question_index = idx.split('-')
            question_index = int(question_index)
            question = questions_by_section[section_name][question_index]
            selected_questions.append(question)

        test_name = request.form.get('test_name')
        test_time = request.form.get('test_time')
    
        html_template = '''
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <title>Test Details</title>
            <link rel="icon" href="{{ url_for('static', filename='images/logo.png') }}" type="image/x-icon">
            <script src="https://polyfill.io/v3/polyfill.min.js?features=es6"></script>
            <script id="MathJax-script" async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>
            <style>
                body {
                    font-family: Arial, sans-serif;
                    display: flex;
                    flex-direction: column;
                    min-height: 100vh;
                }
                .container {
                    flex: 1;
                    width: 95%;
                    margin: 0 auto;
                }
                .header {
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                    padding: 20px;
                    background-color: #f8f9fa;
                    border-bottom: 1px solid #ddd;
                }
                .header img {
                    width: 50px; /* Adjusted size */
                    height: auto;
                }
                .contact-details {
                    text-align: right;
                }
                .test-header {
                    text-align: center;
                    padding: 20px;
                }
                .test-info, .section-info, .question-list {
                    margin-bottom: 20px;
                    padding: 20px;
                    border: 1px solid #ccc;
                    border-radius: 5px;
                }
                .question-list {
                    display: flex;
                    flex-wrap: wrap;
                }
                .question {
                    width: 45%;
                    margin: 2.5%;
                }
                .question p {
                    margin: 5px 0;
                }
                .question-options ol {
                    margin-left: 20px;
                }
                .footer {
                    text-align: center;
                    padding: 10px;
                    background-color: #f8f9fa;
                    border-top: 1px solid #ddd;
                }
            </style>
        </head>
        <body>
            <div class="header">
                <img src="{{ url_for('static', filename='images/logo.png') }}" alt="Logo">
                <div class="contact-details">
                    <p>Contact us: +91 88842 49383</p>
                    <p>Email: admin@kodemapa.com</p>
                </div>
            </div>
            <div class="container">
                <div class="test-header">
                    <h1>{{ test_name }}</h1>
                    <p>Time: {{ test_time }} minutes</p>
                </div>
                <div class="question-list">
                    {% for index, question in enumerate(selected_questions) %}
                        <div class="question" data-qid="{{ question['qid'] }}">
                            <p><strong>Question {{ index + 1 }}:</strong> {{ question['que']['1']['q_string'] | safe }}</p>
                            <div class="question-options">
                                <ol>
                                    {% for option in question['que']['1']['q_option'] %}
                                        <li>{{ option | safe }}</li>
                                    {% endfor %}
                                </ol>
                            </div>
                        </div>
                    {% endfor %}
                </div>
                <button onclick="submitData()">Export Selected Questions to DOCX</button>
                <button type="button" onclick="downloadSolutions('{{ test_name }}')">Download Solutions</button>
            </div>
            <div class="footer">
                <p>Kodemapa Exampad &copy; All rights reserved</p>
            </div>
            <script>
                function submitData() {
                    const testName = "{{ test_name }}";
                    const testTime = "{{ test_time }}";
                    const selectedQuestions = {{ selected_questions | tojson | safe }};
                    
                    const data = {
                        test_name: testName,
                        test_time: testTime,
                        selected_questions: selectedQuestions
                    };

                    fetch('/export_to_docx', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json'
                        },
                        body: JSON.stringify(data)
                    })
                    .then(response => response.blob())
                    .then(blob => {
                        const url = window.URL.createObjectURL(blob);
                        const a = document.createElement('a');
                        a.style.display = 'none';
                        a.href = url;
                        a.download = testName + '.docx';
                        document.body.appendChild(a);
                        a.click();
                        window.URL.revokeObjectURL(url);
                    })
                    .catch(error => console.error('Error:', error));
                }

                function downloadSolutions(testName) {
                    if (!testName) {
                        console.error('Test name is not provided');
                        alert('Test name is missing. Please ensure you have entered a test name.');
                        return;
                    }

                    const selectedQuestions = Array.from(document.querySelectorAll('.question'))
                        .map(questionDiv => questionDiv.getAttribute('data-qid'))
                        .filter(qid => qid); // Filter out any null or undefined values

                    if (selectedQuestions.length === 0) {
                        alert('No questions found. Please ensure questions are loaded before downloading solutions.');
                        return;
                    }

                    const data = {
                        test_name: testName,
                        question_ids: selectedQuestions
                    };

                    fetch('/download_solutions', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json'
                        },
                        body: JSON.stringify(data)
                    })
                    .then(response => {
                        if (!response.ok) {
                            throw new Error('Network response was not ok');
                        }
                        return response.blob();
                    })
                    .then(blob => {
                        const url = window.URL.createObjectURL(blob);
                        const a = document.createElement('a');
                        a.style.display = 'none';
                        a.href = url;
                        a.download = `${testName}_solutions.docx`;
                        document.body.appendChild(a);
                        a.click();
                        window.URL.revokeObjectURL(url);
                    })
                    .catch(error => {
                        console.error('Error:', error);
                        alert('An error occurred while downloading the solutions. Please try again.');
                    });
                }
            </script>
        </body>
        </html>
        '''

        return render_template_string(html_template, selected_questions=selected_questions, enumerate=enumerate, test_name=test_name, test_time=test_time)

    else:
        return render_template('select_questions.html', questions_by_section=questions_by_section, enumerate=enumerate)


def combine_json_files(input_file1, input_file2,input_file3):
    # Load the JSON data from both input files
    with open(input_file1, 'r', encoding='utf-8') as file1:
        data1 = json.load(file1)
        
    with open(input_file2, 'r', encoding='utf-8') as file2:
        data2 = json.load(file2)

    with open(input_file3, 'r', encoding='utf-8') as file2:
        data3 = json.load(file2)
    
    # Combine the data
    combined_data = {**data1, **data3,**data2}
    
    # Return the combined data as a dictionary
    return combined_data

# Example usage
input_file1 = 'Answers/solutions1.json'
input_file2 = 'Answers/solutions2.json'
input_file3 = 'Answers/solutions3.json'

@app.route('/download_solutions', methods=['POST'])
def download_solutions():
    data = request.json
    test_name = data['test_name']
    question_ids = data['question_ids']
    
    # Load the solutions from the JSON file with UTF-8 encoding
    # with open('Answers/solutions.json', 'r', encoding='utf-8') as f:
    #     solutions = json.load(f)
    solutions = combine_json_files(input_file1, input_file2,input_file3)
    # Create a DOCX document
    doc = Document()
    doc.add_heading(f'Solutions for {test_name}', level=1)
    
    # Add solutions to the document
    for index, qid in enumerate(question_ids, start=1):
        solution = solutions.get(qid, {}).get('correct_answers', 'solution not found')[0]
        # solution = solutions.get(str(correct_answers[0]), 'Solution not found')
        # doc.add_heading(f'Question {index}', level=2)
        if isinstance(solution,int) :
            quest = f'[{index}]. {solution+1}'
            doc.add_paragraph(quest)
        elif  isinstance(solution, float):
            quest = f'[{index}]. {solution}'
            doc.add_paragraph(quest)
        elif isinstance(solution,str) :
            q_string = handle_images(solution)
            q_string = handle_tables(q_string)
            q_string = latex_to_markdown(q_string)
            q_string = clean_html(q_string)
            q_string = clean_latex(q_string)
            quest = f'[{index}]. {q_string}'
            # docx_content = pypandoc.convert_text(quest, 'docx', format='md')
            doc.add_paragraph(quest)
        else:
            quest = f'[{index}]. {qid}'
            doc.add_paragraph(quest)
    
    # Save the document
    docx_filename = f'{test_name}_solutions.docx'
    doc.save(docx_filename)
    
    return send_file(docx_filename, as_attachment=True, download_name=f'{test_name}_solutions.docx')

def determine_difficulty(question):
    # Example logic to determine difficulty
    # This should be replaced with actual logic based on your criteria
    if 'easy' in question.lower():
        return 'Easy'
    elif 'moderate' in question.lower():
        return 'Moderate'
    else:
        return 'Difficult'

@app.route('/export_to_docx', methods=['POST'])
def export_to_docx():
    try:
        data = request.get_json(force=True)  # force=True to handle potential mismatched content type
        
        # Pre-process JSON to escape backslashes
        data_str = json.dumps(data)
        # data_str = data_str.replace('\\', '\\\\')
        data = json.loads(data_str)
        
        test_name = data['test_name']
        selected_questions = data['selected_questions']
        
        output_docx_file = 'output.docx'
        convert_json_to_docx(selected_questions, output_docx_file, test_name)

        return send_file(output_docx_file, as_attachment=True, download_name=f'{test_name}.docx')
    except Exception as e:
        print(f"Error in export_to_docx: {str(e)}")
        return jsonify({"error": str(e)}), 500

# Utility Functions

def clean_html(raw_html):
    clean = re.sub(r'<style>.*?</style>', '', raw_html, flags=re.DOTALL)
    clean = clean.replace('<br/>', '\n')
    clean = re.sub(r'<[^>]+>', '', clean)
    return clean.strip()

def clean_latex(text):
    # Remove space after opening $ and before closing $
    text = re.sub(r'\$\s+', '$', text)
    text = re.sub(r'\s+\$', '$', text)
    # Preserve important LaTeX structures
    text = text.replace('\\left', '').replace('\\right', '')
    return text

def process_latex(text):
    # Define a regex pattern for HTML entities
    entity_pattern = r'&[a-zA-Z0-9#]+;'
    
    # Find all HTML entities in the text
    entities = re.findall(entity_pattern, text)
    
    # Unescape each entity individually
    for entity in entities:
        unescaped = html.unescape(entity)
        text = text.replace(entity, unescaped)
    
    # Now process LaTeX-specific patterns
    text = re.sub(r'\\mathrm{([^}]+)}', r'\1', text)  # Remove \mathrm
    text = re.sub(r'\\\(|\\\)', '$', text)  # Replace \( and \) with $
    text = re.sub(r'(\$[^$]+)\$\s*(\$[^$]+\$)', r'\1\2', text)  # Combine adjacent LaTeX expressions
    
    return text

import re

def latex_to_markdown(latex):
    latex = html.unescape(latex)
    # Remove HTML tags
    latex = re.sub(r'</?p>', '', latex)

    # Handle the piecewise function
    piecewise_pattern = r'\\left\\{\\begin\{array\}\{cc\}.*?\\end\{array\}\\right\.'
    piecewise_match = re.search(piecewise_pattern, latex, re.DOTALL)
    
    if piecewise_match:
        piecewise = piecewise_match.group(0)
        # Replace \left\{ and \right. with \{ and \}
        piecewise = piecewise.replace('\\left\\{', '\\{').replace('\\right.', '\\}')
        # Replace \begin{array}{cc} and \end{array} with cases environment
        piecewise = piecewise.replace('\\begin{array}{cc}', '\\begin{cases}').replace('\\end{array}', '\\end{cases}')
        latex = latex.replace(piecewise_match.group(0), piecewise)

    # Replace \( and \) with $
    latex = latex.replace('\\(', '$').replace('\\)', '$')

    # Handle LaTeX symbols and commands
    symbol_map = {
        '\\rightarrow': '→',
        '\\leftarrow': '←',
        '\\Rightarrow': '⇒',
        '\\Leftarrow': '⇐',
        '\\leftrightarrow': '↔',
        '\\Leftrightarrow': '⇔',
        '\\leq': '≤',
        '\\geq': '≥',
        '\\lt': '<',
        '\\gt': '>',
        '\\ne': '≠',
    }
    for latex_symbol, unicode_symbol in symbol_map.items():
        latex = latex.replace(latex_symbol, unicode_symbol)

    return latex.strip()
    
def clean_latex(text):
    # Remove space after opening $ and before closing $
    text = re.sub(r'\$\s+', '$', text)
    text = re.sub(r'\s+\$', '$', text)
    # Preserve important LaTeX structures
    text = text.replace('\\left', '').replace('\\right', '')
    return text

def handle_images(content):
    img_pattern = r'<img[^>]+src="([^">]+)"[^>]*>'
    return re.sub(img_pattern, r'![Image](\1)', content)

def handle_tables(content):
    tables = re.findall(r'<table.*?</table>', content, re.DOTALL)
    for table in tables:
        soup = BeautifulSoup(table, 'html.parser')
        markdown_table = []
        rows = soup.find_all('tr')
        for row in rows:
            cols = row.find_all(['td', 'th'])
            markdown_row = '| ' + ' | '.join([col.get_text(strip=True) for col in cols]) + ' |'
            markdown_table.append(markdown_row)
        markdown_table = '\n'.join(markdown_table)
        content = content.replace(table, markdown_table)
    return content

def json_to_markdown(json_data, test_name):
    questions = json.loads(json_data)
    markdown_content = f"# {test_name}\n\n"
    
    for index, question in enumerate(questions, 1):
        markdown_content += f"## Question {index}\n\n"
        q_string = question['que']['1']['q_string']
        
        q_string = handle_images(q_string)
        q_string = handle_tables(q_string)
        q_string = latex_to_markdown(q_string)
        q_string = clean_html(q_string)
        q_string = clean_latex(q_string)

        markdown_content += f"{q_string}\n\n"
        
        if 'q_option' in question['que']['1']:
            for i, option in enumerate(question['que']['1']['q_option'], 65):
                option_text = option
                option_text = handle_images(option_text)
                option_text = handle_tables(option_text)
                option_text = latex_to_markdown(option_text)
                option_text = clean_html(option_text)
                option_text = clean_latex(option_text)
                markdown_content += f"{chr(i)}) {option_text}\n"
            markdown_content += "\n"
    
    return markdown_content

def set_two_columns(docx_file):
    document = Document(docx_file)
    for section in document.sections:
        sectPr = section._sectPr
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720')
        sectPr.append(cols)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    document.save(docx_file)
    print(f"DOCX file '{docx_file}' has been updated to two columns.")

def convert_json_to_docx(selected_questions, docx_file, test_name):
    try:
        json_data = json.dumps(selected_questions)
        markdown_output = json_to_markdown(json_data, test_name)

        # Print Markdown content to console
        print(markdown_output)
        
        # Convert Markdown to DOCX using pypandoc
        pypandoc.convert_text(markdown_output, 'docx', format='md', outputfile=docx_file)
        
        # Adjust DOCX to two columns
        set_two_columns(docx_file)

        print(f"\nDOCX file '{docx_file}' has been created and updated to two columns successfully.")

    except FileNotFoundError:
        print(f"Error: The file '{selected_questions}' was not found.")
    except json.JSONDecodeError:
        print("Error: The JSON data is not properly formatted.")
    except KeyError as e:
        print(f"Error: The key {e} was not found in the JSON data.")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")


webview.create_window('KodeMapa-Exampad', app)
if __name__ == "__main__":
    app.run(debug=True)
    # webview.start()